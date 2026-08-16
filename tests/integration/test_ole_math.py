"""Integration tests: MathType OLE equation objects -> native OMML.

Builds an in-memory .docx that embeds a synthetic MathType OLE object
(``oleObject1.bin`` with an ``Equation Native`` MTEF stream) and verifies the
engine discovers it, converts it to OMML, and that round-tripping preserves
the change.
"""

from __future__ import annotations

import io

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from tests.unit.rendering.test_mtef_reader import _char, _end, _tmpl, _wrap
from tianshang_scribe.core.word_engine import WordEngine

_NS_O = 'urn:schemas-microsoft-com:office:office'


def _ole_docx(mtef_records: bytes) -> WordEngine:
    """Build a WordEngine whose first paragraph embeds a MathType OLE object."""
    doc = Document()
    p = doc.add_paragraph('eq: ')
    partname = PackURI('/word/embeddings/oleObject1.bin')
    part = Part(
        partname,
        'application/vnd.openxmlformats-officedocument.oleObject',
        bytes(mtef_records),
        doc.part.package,
    )
    r_id = doc.part.relate_to(part, RELATIONSHIP_TYPE.OLE_OBJECT)
    obj = OxmlElement('w:object')
    ole = etree.SubElement(obj, f'{{{_NS_O}}}OLEObject')
    ole.set(qn('r:id'), r_id)
    p._p.append(obj)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    engine = WordEngine()
    import tempfile

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
        tf.write(buf.getvalue())
        tmp = tf.name
    engine.open(tmp)
    import os

    os.unlink(tmp)
    return engine


def _fraction_mtef() -> bytes:
    return _wrap(_tmpl(11) + _char('a') + _char('b') + _end() + _end())


def test_discovers_ole_equation() -> None:
    engine = _ole_docx(_fraction_mtef())
    parts = engine._ole_equation_parts()
    assert len(parts) == 1
    assert parts[0][0].tag == qn('w:object')


def test_extract_math_latex() -> None:
    engine = _ole_docx(_fraction_mtef())
    latex = engine.extract_math_latex()
    assert len(latex) == 1
    assert '\\frac' in latex[0]


def test_convert_ole_equations_replaces_object() -> None:
    engine = _ole_docx(_fraction_mtef())
    converted = engine.convert_ole_equations()
    assert converted == 1
    # The w:object should be gone, replaced by m:oMath.
    xml = engine.doc.paragraphs[0]._p.xml
    assert 'w:object' not in xml
    assert 'm:oMath' in xml


def test_convert_then_set_math_font_persists() -> None:
    engine = _ole_docx(_fraction_mtef())
    assert engine.convert_ole_equations() == 1
    engine.set_math_font('Times New Roman')
    out = io.BytesIO()
    engine.doc.save(out)
    out.seek(0)
    doc2 = Document(out)
    math_font = doc2.settings.element.find(
        './/{http://schemas.openxmlformats.org/officeDocument/2006/math}mathFont'
    )
    assert math_font is not None
    assert (
        math_font.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val')
        == 'Times New Roman'
    )
    assert 'm:oMath' in doc2.paragraphs[0]._p.xml


def test_convert_no_equations_returns_zero() -> None:
    doc = Document()
    doc.add_paragraph('plain text')
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    engine = WordEngine()
    import os
    import tempfile

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
        tf.write(buf.getvalue())
        tmp = tf.name
    engine.open(tmp)
    os.unlink(tmp)
    assert engine.convert_ole_equations() == 0


def test_roundtrip_preserves_conversion() -> None:
    engine = _ole_docx(_fraction_mtef())
    engine.convert_ole_equations()
    io.BytesIO()
    engine.save = lambda p=None: None  # not used
    # Save via python-docx directly and reload.
    out = io.BytesIO()
    engine.doc.save(out)
    out.seek(0)
    doc2 = Document(out)
    assert 'm:oMath' in doc2.paragraphs[0]._p.xml
    assert 'w:object' not in doc2.paragraphs[0]._p.xml


def test_fraction_mtef_roundtrip_value() -> None:
    """The MTEF for \\frac{a}{b} must survive OLE embedding and convert."""
    engine = _ole_docx(_fraction_mtef())
    latex = engine.extract_math_latex()[0]
    # a and b must be present in the emitted LaTeX.
    assert 'a' in latex
    assert 'b' in latex


def test_unknown_ole_part_is_skipped() -> None:
    """A non-MathType OLE object (no Equation Native stream) is skipped."""
    doc = Document()
    p = doc.add_paragraph('obj: ')
    partname = PackURI('/word/embeddings/oleObject2.bin')
    part = Part(
        partname,
        'application/vnd.openxmlformats-officedocument.oleObject',
        b'not-a-cfb',
        doc.part.package,
    )
    r_id = doc.part.relate_to(part, RELATIONSHIP_TYPE.OLE_OBJECT)
    obj = OxmlElement('w:object')
    ole = etree.SubElement(obj, f'{{{_NS_O}}}OLEObject')
    ole.set(qn('r:id'), r_id)
    p._p.append(obj)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    engine = WordEngine()
    import os
    import tempfile

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
        tf.write(buf.getvalue())
        tmp = tf.name
    engine.open(tmp)
    os.unlink(tmp)
    # Invalid CFB must not crash conversion.
    assert engine.convert_ole_equations() == 0
    assert engine.extract_math_latex() == []


def test_add_matheq_object_roundtrip() -> None:
    """Embedding a formula as a MathType object must be re-extractable."""
    engine = WordEngine()
    engine.create()
    engine.add_matheq_object(r'\frac{a}{b} + x^{2}')
    latex = engine.extract_math_latex()
    assert len(latex) == 1
    assert '\\frac' in latex[0]
    assert 'x' in latex[0]


def test_add_matheq_object_persists_and_reopens() -> None:
    import os
    import tempfile

    engine = WordEngine()
    engine.create()
    engine.add_matheq_object(r'\int_0^1 x dx')
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
        tmp = tf.name
    try:
        engine.save(tmp)
        reopened = WordEngine()
        reopened.open(tmp)
        parts = reopened._ole_equation_parts()
        assert len(parts) == 1
        latex = reopened.extract_math_latex()
        assert len(latex) == 1
        assert '\\int' in latex[0]
    finally:
        os.unlink(tmp)


def test_add_matheq_object_then_convert_roundtrip() -> None:
    """A MathType object written by us converts back to OMML cleanly."""
    engine = WordEngine()
    engine.create()
    engine.add_matheq_object(r'\frac{a}{b}')
    converted = engine.convert_ole_equations()
    assert converted == 1
    xml = engine.doc.paragraphs[0]._p.xml
    assert 'm:oMath' in xml
    assert 'w:object' not in xml
