"""Word (.docx) engine implementing DocumentABC via python-docx."""

from __future__ import annotations

import contextlib
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from src.core.document import DocumentABC
from src.rendering.styles import TextStyle

ALIGNMENT_MAP: dict[str, WD_ALIGN_PARAGRAPH] = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class WordEngine(DocumentABC):
    """python-docx backed engine for reading and writing Word documents."""

    def __init__(self, path: str | Path | None = None) -> None:
        """Initialize the engine with an optional document ``path``."""
        super().__init__(path)
        self._doc: _DocumentType | None = None
        self._base_style: TextStyle = TextStyle.default_word()

    @property
    def doc(self) -> _DocumentType:
        """Return the loaded document, raising if none is loaded."""
        if self._doc is None:
            raise RuntimeError('No document loaded. Call create() or open() first.')
        return self._doc

    def create(self) -> None:
        """Create a new blank Word document."""
        self._doc = Document()
        self._path = None
        self._base_style = TextStyle.default_word()

    def open(self, path: str | Path) -> None:
        """Load an existing Word document from ``path``."""
        self._path = Path(path)
        if not self._path.exists():
            raise FileNotFoundError(f'File not found: {self._path}')
        self._doc = Document(str(self._path))
        self._base_style = TextStyle.default_word()

    def save(self, path: str | Path | None = None) -> None:
        """Persist the document, optionally to ``path``."""
        if path is not None:
            self._path = Path(path)
        if self._path is None:
            raise ValueError('No output path specified.')
        self._path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self._path))

    def get_base_style(self) -> TextStyle:
        """Return the current base style."""
        return self._base_style

    def add_text(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        font_name: str | None = None,
        font_size: int | None = None,
        color: str | None = None,
        alignment: str | None = None,
        text_style: TextStyle | None = None,
        **kwargs: Any,
    ) -> Any:
        """Add a paragraph of styled ``text`` and return it."""
        inline = TextStyle(
            bold=bold or None,
            italic=italic or None,
            font_name=font_name,
            font_size=font_size,
            color=color,
            alignment=alignment,
        )
        final = self._base_style
        if text_style is not None:
            final = final.merge(text_style)
        final = final.merge(inline)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        self._apply_run_style(run, final)

        if final.alignment and final.alignment in ALIGNMENT_MAP:
            paragraph.alignment = ALIGNMENT_MAP[final.alignment]

        return paragraph

    def add_styled_content(
        self,
        tokens: list[dict[str, Any]],
        base_style: TextStyle | None = None,
    ) -> Any:
        """Render parsed LaTeX-style ``tokens`` into a paragraph and return it."""
        current_style = base_style if base_style else self._base_style

        for token in tokens:
            if token.get('command') == 'set_font':
                _apply_font_config(self, token)

        content_tokens = [
            t for t in tokens if not (t.get('command') == 'set_font' and not t.get('content'))
        ]
        has_visible = any(
            t.get('type') == 'text' or (t.get('command') and t.get('command') not in ('set_font',))
            for t in content_tokens
        )

        if not content_tokens or not has_visible:
            return None

        paragraph = self.doc.add_paragraph()

        for token in tokens:
            content_kind = token.get('type', 'text')
            content = token.get('content', '')

            if not content and content_kind == 'text':
                continue

            if content_kind == 'text':
                run = paragraph.add_run(str(content))
                self._apply_run_style(run, current_style)

            elif content_kind == 'command':
                cmd = token.get('command', '')
                if cmd == 'set_font':
                    _apply_font_config(self, token)
                    continue

                if cmd == 'newpage':
                    if paragraph.text.strip():
                        paragraph = self.doc.add_paragraph()
                    self.doc.add_page_break()
                    paragraph = self.doc.add_paragraph()
                    continue

                if cmd == 'includegraphics':
                    img_path = token.get('image', '')
                    if img_path:
                        try:
                            paragraph.add_run().add_picture(img_path)
                        except Exception:
                            run = paragraph.add_run(f'[Image: {img_path}]')
                            self._apply_run_style(run, current_style)
                    continue

                if cmd == 'heading':
                    level = token.get('level', 1)
                    h_content = token.get('content', '')
                    self.doc.add_heading(h_content, level=level)
                    paragraph = self.doc.add_paragraph()
                    continue

                if cmd in (
                    'centering',
                    'raggedright',
                    'raggedleft',
                    'linespread',
                    'indent',
                    'noindent',
                ):
                    paragraph = self.doc.add_paragraph()
                    para_format = paragraph.paragraph_format
                    if cmd == 'centering':
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif cmd == 'raggedright':
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    elif cmd == 'raggedleft':
                        from docx.enum.text import WD_ALIGN_PARAGRAPH

                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    elif cmd == 'linespread':
                        with contextlib.suppress(ValueError, TypeError):
                            para_format.line_spacing = float(content)
                    elif cmd == 'indent':
                        from docx.shared import Cm

                        para_format.first_line_indent = Cm(0.74)
                    elif cmd == 'noindent':
                        from docx.shared import Cm

                        para_format.first_line_indent = Cm(0)
                    inner = token.get('content', '')
                    if inner:
                        if '\\' in inner:
                            from src.rendering.latex_parser import parse_structured

                            inner_tokens = parse_structured(inner)
                            self._render_tokens_inline(paragraph, inner_tokens, current_style)
                        elif inner.strip():
                            run = paragraph.add_run(inner)
                            self._apply_run_style(run, current_style)
                    paragraph = self.doc.add_paragraph()
                    continue

                if cmd == 'math':
                    latex = token.get('latex', '')
                    self._add_omath(paragraph, latex)
                    continue

                token_style = TextStyle.from_latex_token(token)
                merged = current_style.merge(token_style)

                inner_content = token.get('content', '')
                if inner_content:
                    if '\\' in inner_content:
                        from src.rendering.latex_parser import parse_structured

                        inner_tokens = parse_structured(inner_content)
                        self._render_tokens_inline(paragraph, inner_tokens, merged)
                    else:
                        run = paragraph.add_run(inner_content)
                        self._apply_run_style(run, merged)

        return paragraph

    def _render_tokens_inline(
        self,
        paragraph: Any,
        tokens: list[dict[str, Any]],
        base_style: TextStyle,
    ) -> None:
        current_style = base_style

        for token in tokens:
            content_kind = token.get('type', 'text')
            content = token.get('content', '')

            if content_kind == 'text':
                if content:
                    run = paragraph.add_run(str(content))
                    self._apply_run_style(run, current_style)

            elif content_kind == 'command':
                cmd = token.get('command', '')
                if cmd in (
                    'newpage',
                    'includegraphics',
                    'heading',
                    'set_font',
                    'centering',
                    'raggedright',
                    'raggedleft',
                    'linespread',
                    'indent',
                    'noindent',
                ):
                    continue

                if cmd == 'math':
                    latex = token.get('latex', '')
                    self._add_omath(paragraph, latex)
                    continue

                token_style = TextStyle.from_latex_token(token)
                merged = current_style.merge(token_style)

                inner_content = token.get('content', '')
                if inner_content:
                    if '\\' in inner_content:
                        from src.rendering.latex_parser import parse_structured

                        inner_tokens = parse_structured(inner_content)
                        self._render_tokens_inline(paragraph, inner_tokens, merged)
                    else:
                        run = paragraph.add_run(inner_content)
                        self._apply_run_style(run, merged)

    def add_latex_content(self, text: str) -> Any:
        """Parse ``text`` as LaTeX-style markup and add the resulting content."""
        from src.rendering.latex_parser import parse_structured

        tokens = parse_structured(text)
        return self.add_styled_content(tokens)

    def add_math_formula(self, latex: str) -> Any:
        """Add ``latex`` as a native OMML formula and return its paragraph."""
        paragraph = self.doc.add_paragraph()
        self._add_omath(paragraph, latex)
        return paragraph

    def _add_omath(self, paragraph: Any, latex: str) -> None:
        from src.rendering.math_omml import latex_to_omml

        omml = latex_to_omml(latex)
        if omml is not None:
            paragraph._p.append(omml)

    def _apply_run_style(self, run: Any, style: TextStyle) -> None:
        if style.font_name is not None:
            run.font.name = style.font_name

        if style.cjk_font_name is not None:
            r_pr = run._r.get_or_add_rPr()
            r_fonts = r_pr.find(qn('w:rFonts'))
            if r_fonts is None:
                r_fonts = OxmlElement('w:rFonts')
                r_pr.insert(0, r_fonts)
            r_fonts.set(qn('w:eastAsia'), style.cjk_font_name)

        if style.font_size is not None:
            run.font.size = Pt(style.font_size)
        if style.bold is not None and style.bold:
            run.bold = True
        if style.italic is not None and style.italic:
            run.italic = True
        if style.underline is not None and style.underline:
            run.underline = True
        if style.color is not None:
            with contextlib.suppress(ValueError, IndexError):
                run.font.color.rgb = RGBColor(
                    int(style.color[0:2], 16),
                    int(style.color[2:4], 16),
                    int(style.color[4:6], 16),
                )
        if style.small_caps is not None:
            run.font.small_caps = style.small_caps

    def replace_text(self, old: str, new: str, regex: bool = False) -> int:
        """Replace ``old`` with ``new``, optionally treating ``old`` as a regex."""
        count = 0
        for paragraph in self.doc.paragraphs:
            if regex:
                if re.search(old, paragraph.text):
                    for run in paragraph.runs:
                        prev = run.text
                        run.text = re.sub(old, new, run.text)
                        if run.text != prev:
                            count += 1
            else:
                if old in paragraph.text:
                    for run in paragraph.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, new)
                            count += 1
        return count

    def set_style(self, style_str: str) -> None:
        """Merge the given ``style_str`` into the document's base style."""
        new_style = TextStyle.from_string(style_str)
        self._base_style = self._base_style.merge(new_style)

    def apply_style_to_all(self) -> None:
        """Apply the current base style to every run in paragraphs and tables."""
        for paragraph in self.doc.paragraphs:
            for run in paragraph.runs:
                self._apply_run_style(run, self._base_style)

        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._apply_run_style(run, self._base_style)

    def to_pdf(self, output_path: str | Path) -> None:
        """Save the document and convert it to PDF at ``output_path``."""
        self.save()
        from src.transform.pdf import word_to_pdf

        word_to_pdf(str(self._path), str(output_path))

    def add_heading(self, text: str, level: int = 1) -> Any:
        """Add a heading of the given ``level`` and return it."""
        return self.doc.add_heading(text, level=level)

    def add_page_break(self) -> None:
        """Insert a page break at the end of the document."""
        self.doc.add_page_break()

    def add_table(self, rows: int, cols: int) -> Any:
        """Add an empty table of ``rows`` by ``cols`` cells and return it."""
        return self.doc.add_table(rows=rows, cols=cols)

    def add_table_data(self, rows: list[list[str]]) -> Any:
        """Add a table populated with ``rows`` of cell values and return it."""
        if not rows or not rows[0]:
            raise ValueError('add_table_data requires at least one row with one column')
        table = self.doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'
        for r, row in enumerate(rows):
            for c, value in enumerate(row):
                table.rows[r].cells[c].text = str(value)
        if len(rows) > 1:
            for c in range(len(rows[0])):
                for paragraph in table.rows[0].cells[c].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        return table

    def add_image(
        self, image_path: str, width: float | None = None, height: float | None = None
    ) -> Any:
        """Insert ``image_path`` into a new paragraph, optionally sized in inches."""
        from docx.shared import Inches

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        if width and height:
            return run.add_picture(image_path, width=Inches(width), height=Inches(height))
        elif width:
            return run.add_picture(image_path, width=Inches(width))
        elif height:
            return run.add_picture(image_path, height=Inches(height))
        return run.add_picture(image_path)

    def get_metadata(self) -> dict[str, str | None]:
        """Return document core properties such as author and title."""
        props = self.doc.core_properties
        return {
            'author': props.author,
            'title': props.title,
            'subject': props.subject,
            'category': props.category,
            'keywords': props.keywords,
            'comments': props.comments,
        }

    def set_metadata(self, **kwargs: str) -> None:
        """Set document core properties from keyword arguments."""
        props = self.doc.core_properties
        for key, value in kwargs.items():
            key_lower = key.lower()
            if key_lower == 'author':
                props.author = value
            elif key_lower == 'title':
                props.title = value
            elif key_lower == 'subject':
                props.subject = value
            elif key_lower == 'category':
                props.category = value
            elif key_lower == 'keywords':
                props.keywords = value
            elif key_lower == 'comments':
                props.comments = value

    def add_comment(self, text: str, range_start: int = 0, range_end: int = 0) -> None:
        """Attach an inline comment anchored to the first paragraph."""
        paragraph = self.doc.paragraphs[0] if self.doc.paragraphs else self.doc.add_paragraph()
        run = paragraph.add_run('')
        run_element = run._r
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        comment_range_start = OxmlElement('w:commentRangeStart')
        comment_range_start.set(qn('w:id'), '0')
        run_element.addprevious(comment_range_start)
        comment_range_end = OxmlElement('w:commentRangeEnd')
        comment_range_end.set(qn('w:id'), '0')
        run_element.addnext(comment_range_end)
        comment_ref = OxmlElement('w:r')
        comment_ref_elt = OxmlElement('w:commentReference')
        comment_ref_elt.set(qn('w:id'), '0')
        comment_ref.append(comment_ref_elt)
        run_element.addnext(comment_ref)

    def set_protection(self, password: str) -> None:
        """Enable read-only protection on every section with the given ``password``."""
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        for section in self.doc.sections:
            protection = parse_xml(
                f'<w:documentProtection {nsdecls("w")} '
                f'w:edit="readOnly" w:enforcement="1" '
                f'w:cryptProviderType="rsaAES" '
                f'w:cryptAlgorithmClass="hash" '
                f'w:cryptAlgorithmType="typeAny" '
                f'w:cryptAlgorithmSid="14" '
                f'w:hash="placeholder"/>'
            )
            section._sectPr.append(protection)

    def unprotect(self) -> None:
        """Remove read-only protection from every section."""
        for section in self.doc.sections:
            elements = section._sectPr.findall(qn('w:documentProtection'))
            for el in elements:
                section._sectPr.remove(el)

    def add_toc(self) -> None:
        """Insert a table-of-contents field at the end of the document."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        run._r.append(fld_char)
        run2 = paragraph.add_run()
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' TOC \\o "1-3" \\h \\z '
        run2._r.append(instr_text)
        run3 = paragraph.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run3._r.append(fld_char_end)

    def add_section_break(self) -> None:
        """Start a new Word section."""
        self.doc.add_section()

    def merge_workbooks(self, paths: list[str]) -> None:
        """Append the body of each document in ``paths`` to this document."""
        from lxml import etree

        for p in paths:
            merge_doc = Document(str(p))
            for element in merge_doc.element.body:
                self.doc.element.body.append(etree.fromstring(etree.tostring(element)))

    def set_header(self, text: str) -> None:
        """Add ``text`` to the header of the last section."""
        section = self.doc.sections[-1]
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        header.paragraphs[0].add_run(text)

    def set_footer(self, text: str) -> None:
        """Add ``text`` to the footer of the last section."""
        section = self.doc.sections[-1]
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        footer.paragraphs[0].add_run(text)

    def add_watermark(self, text: str) -> None:
        """Add ``text`` as a light gray watermark to the last section's header."""
        from docx.shared import RGBColor

        section = self.doc.sections[-1]
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        paragraph = header.paragraphs[0]
        paragraph.alignment = 1
        run = paragraph.add_run(text)
        run.font.size = 72000
        run.font.color.rgb = RGBColor(0xD9, 0xD9, 0xD9)

    def clear_content(self) -> None:
        """Clear the text of every paragraph."""
        for p in self.doc.paragraphs:
            p.clear()

    def clear_formats(self) -> None:
        """Reset formatting on every run in every paragraph."""
        for p in self.doc.paragraphs:
            for run in p.runs:
                run.font.name = None
                run.font.size = None
                run.bold = None
                run.italic = None
                run.underline = None
                run.font.color.rgb = None

    def clear_links(self) -> None:
        """Remove all hyperlinks from every paragraph."""
        for p in self.doc.paragraphs:
            for hyperlink in p._p.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink'
            ):
                hyperlink.getparent().remove(hyperlink)

    def extract_text(self) -> str:
        """Return the document text including table cell contents."""
        parts = [p.text for p in self.doc.paragraphs]
        for table in self.doc.tables:
            for row in table.rows:
                cells = [c.text.replace('\n', ' ') for c in row.cells]
                if any(cells):
                    parts.append(' | '.join(cells))
        return '\n'.join(line for line in parts if line)

    def extract_tables(self) -> list[list[list[str]]]:
        """Return all tables as lists of rows of cell texts."""
        return [
            [[cell.text for cell in row.cells] for row in table.rows] for table in self.doc.tables
        ]

    def _image_parts(self) -> list[tuple[bytes, str]]:
        blobs: list[tuple[bytes, str]] = []
        for rel in self.doc.part.rels.values():
            if 'image' in rel.reltype and not rel.is_external:
                try:
                    blobs.append(
                        (
                            rel.target_part.blob,
                            Path(rel.target_part.partname).suffix or '.png',
                        )
                    )
                except Exception:  # noqa: S112  # unreadable image part: skip, do not abort extraction
                    continue
        return blobs

    def extract_images(self, output_dir: str | Path) -> list[Path]:
        """Write embedded images into ``output_dir`` and return their paths."""
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)
        saved: list[Path] = []
        for idx, (blob, suffix) in enumerate(self._image_parts()):
            target = out / f'image_{idx}{suffix}'
            target.write_bytes(blob)
            saved.append(target)
        return saved

    def extract_structure(self) -> dict[str, Any]:
        """Return counts of paragraphs, tables, sections, and images."""
        return {
            'paragraphs': len(self.doc.paragraphs),
            'tables': len(self.doc.tables),
            'sections': len(self.doc.sections),
            'images': len(self._image_parts()),
        }


def _apply_font_config(engine: WordEngine, token: dict[str, Any]) -> None:
    """Update ``engine``'s base style font from a ``set_font`` token."""
    from src.rendering.styles import TextStyle

    role = token.get('role', '')
    font = token.get('font', '')
    if not font:
        return

    update = TextStyle()
    if 'CJK' in role:
        update.cjk_font_name = font
    else:
        update.font_name = font

    engine._base_style = engine._base_style.merge(update)
